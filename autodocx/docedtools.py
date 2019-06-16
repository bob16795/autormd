from pathlib import Path
from autodocx.nodes import *
import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from autodocx.formats import list_number
from autodocx.docedequation import *

def get_sub_toc(file, l, sec_num):
    with open(file, encoding='utf-8') as file:
        file_cached = []
        for i in file:
            file_cached.append(i)
    file_cached = "".join(file_cached).replace("\\", "")
    file_cached = file_cached.split("\n")
    #for i in range(len(file_cached)):
    #    file_cached[i] = file_cached[i].replace("\n", " ")
    z = 0
    y = 0
    for i in file_cached:
        if i[:6] == "##### ":
            x += 1
            l.append(f"\t{i[6:]}\t{sec_num}{z}.{y}.{x}")
        elif i[:5] == "#### ":
            x = 0
            y += 1
            l.append(f"\t{i[5:]}\t{sec_num}{z}.{y}")
        elif i[:4] == "### ":
            x = 0
            y = 0
            z += 1
            l.append(f"\t{i[4:]}\t{sec_num}{z}")
    return l

def setcol(doc, num):
    section = doc.add_section(0)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),str(num))

def add_pagenum(doc, section):
    doc.is_linked_to_previous = False
    paragraph = doc.add_paragraph(f"{section}\t", "header top")
    paragraph.paragraph_format.tab_stops.add_tab_stop(Inches(4.5), 2)
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'PAGE \\*Arabic \\* MERGEFORMAT'   # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = ""
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

def add_toc_section(doc):
    section = doc.sections[-1]
    add_pagenum(section.header, "Table of Contents")
    #sectPr = section._sectPr
    #cols = sectPr.xpath('./w:cols')[0]
    #cols.set(qn('w:num'),'2')

    paragraph = doc.paragraphs[-1]
    paragraph.text = "Table of Contents"
    paragraph.style = "Heading TOC"
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'   # change 1-3 depending on heading levels you need
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = ""
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)

def add_index_section(doc):
    section = doc.add_section(4)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    add_pagenum(section.header, "Index")

    doc.add_paragraph("Index", 'Heading 1')
    section = doc.add_section(0)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')

    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    fldChar = OxmlElement('w:fldChar')  # creates a new element
    fldChar.set(qn('w:fldCharType'), 'begin')  # sets attribute on element
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')  # sets attribute on element
    instrText.text = 'INDEX \\h "A" \\c "3" \\z "1033"'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:t')
    fldChar3.text = ""
    fldChar2.append(fldChar3)

    fldChar4 = OxmlElement('w:fldChar')
    fldChar4.set(qn('w:fldCharType'), 'end')

    r_element = run._r
    r_element.append(fldChar)
    r_element.append(instrText)
    r_element.append(fldChar2)
    r_element.append(fldChar4)


def MarkIndexEntry(entry,paragraph):
    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'begin')
    r.append(fldChar)

    run = paragraph.add_run()
    r = run._r
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' XE "%s" '%(entry)
    r.append(instrText)

    run = paragraph.add_run()
    r = run._r
    fldChar = OxmlElement('w:fldChar')
    fldChar.set(qn('w:fldCharType'), 'end')
    r.append(fldChar)

class Token():
    def __init__(self, type, value, at):
        self.type = type
        self.at = at
        self.value = value
    def __str__(self):
        return(f"<type: {self.type}, value: {self.value}>")

class Token_List():
    def __init__(self, tokens):
        self.tokens = tokens

    def length(self):
        return len(self.tokens)

    def peek_or(self, choices):
        for i in choices:
            if self.peek(i):
                return True
        return False

    def peek(self, types):
        peeked = [i.type for i in self.tokens[:(len(types))]]
        if peeked == types:
            return True
        return False

    def peek_at(self, index, types):
        return self.offset(index).peek(types)
    
    def offset(self, index):
        if index == 0:
            return self
        return Token_List(self.tokens[index:])
    
    def grab(self, index):
        return self.tokens[index]

def tokenize(file_cached):
    token_dict = {
        "_" : ["UNDERSCORE", ""],
        "`" : ["GRAVE", ""],
        "*" : ["STAR", ""],
        "+" : ["PLUS", ""],
        "-" : ["MINUS", ""],
        "#" : ["HASH", ""],
        "<" : ["TAGO", ""],
        ">" : ["TAGC", ""],
        "\n": ["NEWLINE", ""],
    }
    tokens = []
    ignore = 0
    for x, i in enumerate(file_cached):
        text = ""
        l = i.strip(" ")
        for y, j in enumerate(l):
            if ignore != 0:
                ignore -= 1
                text = f"{text}{j}"
            else:
                if j == "\\":
                    ignore = 1
                elif j in token_dict:
                    if text != "":
                        tokens.append(Token("TEXT", text, f"line {x}, col {y}"))
                        text = ""
                    token = token_dict[j]
                    tokens.append(Token(token[0], token[1], f"line {x}, col {y}"))
                else:
                    text = f"{text}{j}"
        if text != "":
            tokens.append(Token("TEXT", text, f"line {x}, col {y}"))
        tokens.append(Token("NEWLINE", "", f"line {x}, col {y}"))
    tokens.append(Token("EOF", "", f"line {x}, col {y}"))
    return Token_List(tokens)

def match_star(tokens, With):
    matched_nodes = []
    consumed = 0
    parser = With

    while consumed < tokens.length():
        node = parser(tokens.offset(consumed))
        if type(node) == nullNode:
            break
        matched_nodes.append(node)
        consumed += node.consumed
    return matched_nodes, consumed
    
def match_stars(tokens, Withs):
    matched_nodes = []
    consumed = 0
    parsers = Withs

    while consumed < tokens.length():
        node = match_first(tokens.offset(consumed), parsers)
        if type(node) == nullNode:
            break
        matched_nodes.append(node)
        consumed += node.consumed
    return matched_nodes, consumed
 
def match_first(tokens, matches):
    for j in matches:
        node = j(tokens)
        if type(node) != nullNode:
            return node
    return nullNode()

class parser():
    def Parse_Text(tokens):
        if tokens.peek(["TEXT"]):
            return Node("TEXT", tokens.grab(0).value, 1)
        return nullNode()
    def single_newline(tokens):
        node = nullNode()
        if tokens.peek(["NEWLINE"]):
            node = Node("TEXT", " ", 1)
        if tokens.peek(["NEWLINE", "NEWLINE"]):
            return nullNode()
        return node
    def Parse_Bold(tokens):
        if tokens.peek_or([["STAR", "STAR", "TEXT", "STAR", "STAR"], ["UNDERSCORE", "UNDERSCORE", "TEXT", "UNDERSCORE", "UNDERSCORE"]]):
            return Node("BOLD", tokens.grab(2).value, 5)
        return nullNode()
    def Parse_Emph(tokens):
        if tokens.peek_or([["STAR", "TEXT", "STAR"], ["UNDERSCORE", "TEXT", "UNDERSCORE"]]):
            return Node("EMPH", tokens.grab(1).value, 3)
        return nullNode()
    def Parse_Tag(tokens):
        if tokens.peek(["TAGO", "TEXT", "TAGC"]):
            return Node("TAG", tokens.grab(1).value, 3)
        return nullNode()
    def Parse_Tags(tokens):
        node = match_first(tokens, [parser.single_newline, parser.Parse_Tag])
        if type(node) == nullNode:
            return nullNode()
        if not tokens.peek_at(node.consumed, ['NEWLINE', 'NEWLINE']):
            return nullNode()
        node.consumed += 2
        return node
    def H1_Parser(tokens):
        if tokens.peek(["HASH", "TEXT"]):
            return HeadNode("HEAD1", tokens.grab(1).value, 2)
        return nullNode()
    def H2_Parser(tokens):
        if tokens.peek(["HASH", "HASH", "TEXT"]):
            return HeadNode("HEAD2", tokens.grab(2).value, 3)
        return nullNode()
    def H3_Parser(tokens):
        if tokens.peek(["HASH", "HASH", "HASH", "TEXT"]):
            return HeadNode("HEAD3", tokens.grab(3).value, 4)
        return nullNode()
    def Header_Parser(tokens):
        node = match_first(tokens, [parser.H1_Parser, parser.H2_Parser, parser.H3_Parser])
        if type(node) == nullNode:
            return nullNode()
        if not tokens.peek_at(node.consumed, ['NEWLINE', 'NEWLINE']):
            return nullNode()
        node.consumed += 2
        return node
    def L1_Parser(tokens):
        if tokens.peek(["STAR"]):
            return Node("LIST1", "", 1)
        return nullNode()
    def L2_Parser(tokens):
        if tokens.peek(["PLUS"]):
            return Node("LIST2", "", 1)
        return nullNode()
    def L3_Parser(tokens):
        if tokens.peek(["MINUS"]):
            return Node("LIST3", "", 1)
        return nullNode()
    def Item_Parser(tokens):
        return match_first(tokens, [parser.Sentence_Parser, parser.L1_Parser, parser.L2_Parser, parser.L3_Parser])
    def List_Parser(tokens):
        nodes, consumed = match_star(tokens, parser.Item_Parser)
        if nodes == []:
            return nullNode()
        if not tokens.peek_at(consumed, ['NEWLINE', 'NEWLINE']):
            return nullNode()
        consumed += 2
        return ListNode(nodes, consumed)
    def Code_Start_Parser(tokens):
        if tokens.peek(["GRAVE","GRAVE","GRAVE"]):
            return Node("CODE", "", 3)
        return nullNode()
    def Code_Parser(tokens):
        return match_first(tokens, [parser.Code_Start_Parser, parser.single_newline, parser.Sentence_Parser])
    def Code_Parser_Multiline(tokens):
        nodes, consumed = match_star(tokens, parser.Code_Parser)
        if not tokens.peek(["GRAVE","GRAVE","GRAVE"]):
            return nullNode()
        if nodes == []:
            return nullNode()
        if not tokens.peek_at(consumed, ['NEWLINE', 'NEWLINE']):
            return nullNode()
        consumed += 2
        return CodeNode(nodes, consumed)
    def Sentence_Parser(tokens):
        return match_first(tokens, [parser.Parse_Text, parser.Parse_Emph, parser.Parse_Bold, parser.single_newline])
    def Sentences_NL_Parser(tokens):
        nodes, consumed = match_star(tokens, parser.Sentence_Parser)
        if nodes == []:
            return nullNode()
        if not tokens.peek_at(consumed, ['NEWLINE', 'NEWLINE']):
            return nullNode()
        consumed += 2
        return ParagraphNode(nodes, consumed)
    def Sentences_EOF_Parser(tokens):
        nodes, consumed = match_star(tokens, parser.Sentence_Parser)
        if nodes == []:
            return nullNode()
        if tokens.peek_at(consumed, ['EOF']):
            consumed += 1
        elif tokens.peek_at(consumed, ['NEWLINE', 'EOF']):
            consumed += 2
        else:
            return nullNode()
        return ParagraphNode(nodes, consumed)
    def Parse_Paragraph(tokens):
        return match_first(tokens, [parser.Parse_Tags, parser.Sentences_NL_Parser, parser.Sentences_EOF_Parser, parser.Code_Parser_Multiline, parser.Header_Parser, parser.List_Parser])
    def Parse_Body(tokens):
        nodes, consumed = match_star(tokens, parser.Parse_Paragraph)
        if nodes == []:
            return nullNode()
        return BodyNode(nodes, consumed)

def parse(tokens, file):
    body = parser.Parse_Body(tokens)
    if body.consumed == 1 + tokens.length():
        print(f"error in {file.name} at {tokens.grab(body.consumed).at}")
    return body

def add_to_doc(doc, parsed, file, main):
    if type(parsed) is not nullNode:
        for i in parsed.paragraphs:
            if type(i) is HeadNode or type(i) is Node or type(i) is nullNode:
                if main:
                    if i.type == "HEAD1":
                        doc.add_paragraph(i.value, style="Heading 3")
                    elif i.type == "HEAD2":
                        doc.add_paragraph(i.value, style="Heading 4")
                    elif i.type == "HEAD3":
                        doc.add_paragraph(i.value, style="Heading 5")
                else:
                    if i.type == "HEAD1":
                        doc.add_paragraph(i.value, style="Heading 1")
                    elif i.type == "HEAD2":
                        doc.add_paragraph(i.value, style="Heading 2")
                    elif i.type == "HEAD3":
                        doc.add_paragraph(i.value, style="Heading 3")
                if i.type == "TAG":
                    if i.value == "3COL":
                        setcol(doc, 3)
                    elif i.value == "1COL":
                        setcol(doc, 1)
                    elif i.value == "2COL":
                        setcol(doc, 2)
            elif type(i) is ListNode:
                p = None
                for j in i.sentences:
                    if j.type == "LIST1":
                        p = doc.add_paragraph(style = "List Bullet")
                    elif j.type == "LIST2":
                        p = doc.add_paragraph(style = "List Bullet 2")
                    elif j.type == "LIST3":
                        p = doc.add_paragraph(style = "List Bullet 3")
                    elif j.value != "" and j.value != " " and j.type == "TEXT":
                        if p is None:
                            p = doc.add_paragraph(style = "List Bullet")
                        p.add_run(f"{j.value} ")
                    elif j.value != "" and j.value != " " and j.type == "EMPH":
                        r = p.add_run(f"{j.value} ")
                        r.italic = True
                    elif j.value != "" and j.value != " " and j.type == "BOLD":
                        r = p.add_run(f"{j.value} ")
                        r.bold = True
            elif type(i) is CodeNode:
                p = None
                for j in i.sentences:
                    if j.value != "" and j.value != " " and j.type == "TEXT":
                        p = doc.add_paragraph(j.value, style = "Source Code")
                    elif j.value != "" and j.value != " " and j.type == "EMPH":
                        p = doc.add_paragraph(j.value, style = "Source Code")
                        p.italic = True
                    elif j.value != "" and j.value != " " and j.type == "BOLD":
                        p = doc.add_paragraph(j.value, style = "Source Code")
                        p.bold = True
            else:
                p = doc.add_paragraph(style = "Normal")
                for j in i.sentences:
                    if j.type == "TEXT":
                        p.add_run(f"{j.value}")
                    elif j.value != "" and j.value != " " and j.type == "EMPH":
                        r = p.add_run(f"{j.value}")
                        r.italic = True
                    elif j.value != "" and j.value != " " and j.type == "BOLD":
                        r = p.add_run(f"{j.value}")
                        r.bold = True
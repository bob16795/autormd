import inspect
import os
import docx
from pathlib import Path

import win32com.client

from autodocx.docedtools import *


def add_to_main(doc, file, title, idx):
    p = doc.add_paragraph(title, "Heading 2")
    for i in idx:
        MarkIndexEntry(i, p)
    with open(file, encoding='utf-8') as file:
        file_cached = ""
        for i in file:
            file_cached = f"{file_cached}{i}"
    file_cached = file_cached.split("\n")
    tokens = tokenize(file_cached)
    parsed = parse(tokens, file)
    add_to_doc(doc, parsed, file, True)

def create_main(docdir, cfgdir):
    Docdir = Path(docdir)
    Cfgdir = Path(cfgdir)
    Pdfdir = Docdir / "pdf"
    doc = docx.Document(Cfgdir / "main.docx")

    doc.styles.add_style("Quote", 1, True)

    add_toc_section(doc)

    doc.save(f'{Pdfdir}/main.docx')

def render_main(linux, docdir, cfgdir):
    Docdir = Path(docdir)
    Cfgdir = Path(cfgdir)
    Srcdir = Docdir / "src"
    create_main(docdir, cfgdir)

    with open(Cfgdir / "sections", "r", encoding='utf-8') as file:
        sections = file.read()[:-1].split("\n")

    with open(Cfgdir / "index.csv", "r", encoding='utf-8') as file:
        read = file.read()[:-1].split("\n")
        index = {}
        for i in read:
            data = i.split(",")
            index[data[0]] = data[1:]

    doc = docx.Document(f'{docdir}/pdf/main.docx')
    print(" + pass 1")
    for y, section in enumerate(sorted(sections)):
        print(f" +-- {section}")
        s = doc.add_section(4)
        sectPr = s._sectPr
        add_pagenum(s.header, section)
        sectPr = s._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'1')
        doc.add_paragraph(f" {section}", "Heading 1")
        s = doc.add_section(0)
        sectPr = s._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'2')
        for z, file in enumerate(sorted(Path(Srcdir).glob(f"{section}_*.doc"))):
            title = file.name.replace("The_", "")\
                    .replace(f"{section}_", "")\
                    .replace(f".doc", "")\
                    .replace(f"_", " ")
            queue = [f"{title}\t{y}.{z}"]
            queue = get_sub_toc(file, queue, f"{y}.{z}.")
            for q in queue:
                p = doc.add_paragraph(q, "Chapter toc")
                p.paragraph_format.tab_stops.add_tab_stop(Inches(.125))
                p.paragraph_format.tab_stops.add_tab_stop(Inches(2), 2)
        s = doc.add_section(2)
        sectPr = s._sectPr
        cols = sectPr.xpath('./w:cols')[0]
        cols.set(qn('w:num'),'1')

        for z, file in enumerate(sorted(Path(Srcdir).glob(f"{section}_*.doc"))):
            title = file.name.replace("The_", "")\
                    .replace(f"{section}_", "")\
                    .replace(f".doc", "")\
                    .replace(f"_", " ")
            add_to_main(doc, file, f" {title}", index[file.name])
    add_index_section(doc)
    print(" + saving")
    doc.save(f'{docdir}/pdf/main.docx')
    if not linux:
        update_toc(f'{docdir}/pdf/main.docx')

def update_toc(docx_file):
    word = win32com.client.DispatchEx("Word.Application")
    doc = word.Documents.Open(docx_file)
    selection = word.Selection
    selection.WholeStory()
    selection.Fields.Update()
    doc.Close(SaveChanges=True)
    word.Quit()

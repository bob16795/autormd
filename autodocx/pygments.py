from pygments.formatter import Formatter
from docx.shared import RGBColor


class Word(Formatter):

    def __init__(self, **options):
        Formatter.__init__(self, **options)

        # create a dict of (start, end) tuples that wrap the
        # value of a token so that we can use it in the format
        # method later
        self.styles = {}

        # we iterate over the `_styles` attribute of a style item
        # that contains the parsed style values.
        for token, style in self.style:
            start = end = ''
            # a style item is a tuple in the following form:
            # colors are readily specified in hex: 'RRGGBB'
            if style['color'] == None:
                color = RGBColor.from_string("000000")
            else:
                color = RGBColor.from_string(style['color'])
            self.styles[token] = (color, style['bold'],
                                  style['italic'], style['underline'])

    def format(self, tokensource, outfile):
        # lastval is a string we use for caching
        # because it's possible that an lexer yields a number
        # of consecutive tokens with the same token type.
        # to minimize the size of the generated html markup we
        # try to join the values of same-type tokens here
        lastval = ''
        lasttype = None
        doc = outfile
        self.p = doc.add_paragraph(style="Source Code")
        # wrap the whole output with <pre>

        for ttype, value in tokensource:
            # if the token type doesn't exist in the stylemap
            # we try it with the parent of the token type
            # eg: parent of Token.Literal.String.Double is
            # Token.Literal.String
            while ttype not in self.styles:
                ttype = ttype.parent
            if ttype == lasttype:
                # the current token type is the same of the last
                # iteration. cache it
                lastval += value
            else:
                # not the same token as last iteration, but we
                # have some data in the buffer. wrap it with the
                # defined style and write it to the output file
                if lastval:
                    color, bold, italic, underline = self.styles[lasttype]
                    r = self.p.add_run(lastval)
                    r.font.color.rgb = color
                    r.font.bold = bold
                    r.font.italic = italic
                    r.font.underline = underline
                # set lastval/lasttype to current5 values
                lastval = value
                lasttype = ttype

        # if something is left in the buffer, write it to the
        if lastval:
            color, bold, italic, underline = self.styles[lasttype]
            r = self.p.add_run(lastval.replace("\n", ""))
            r.font.color.rgb = color
            r.font.bold = bold
            r.font.italic = italic
            r.font.underline = underline
